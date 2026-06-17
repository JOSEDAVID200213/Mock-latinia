"""
Pipeline de extracción de texto multi-formato.
Soporta TXT, MD, DOCX, PDF, RTF con evaluación de calidad.
"""
from __future__ import annotations

import logging
import mimetypes
import re
from pathlib import Path

from app.models.schemas import ExtractionResult, FileFormat

logger = logging.getLogger(__name__)


class ExtractionError(Exception):
    """Error durante la extracción de texto."""
    pass


class FileExtractor:
    """
    Extrae texto de diferentes formatos de archivo.
    Evalúa la calidad del texto extraído para adaptar el prompt del LLM.
    """

    EXTENSION_MAP = {
        ".txt": FileFormat.TXT,
        ".md": FileFormat.MD,
        ".docx": FileFormat.DOCX,
        ".pdf": FileFormat.PDF,
        ".rtf": FileFormat.RTF,
    }

    def extract(self, file_path: Path, original_filename: str) -> ExtractionResult:
        """
        Extrae texto del archivo y evalúa su calidad.
        
        Args:
            file_path: Ruta al archivo temporal subido
            original_filename: Nombre original del archivo
            
        Returns:
            ExtractionResult con texto, metadata y score de calidad
        """
        extension = Path(original_filename).suffix.lower()
        file_format = self.EXTENSION_MAP.get(extension)

        if not file_format:
            raise ExtractionError(
                f"Formato no soportado: '{extension}'. "
                f"Formatos válidos: {', '.join(self.EXTENSION_MAP.keys())}"
            )

        # Llamar al parser específico
        extract_method = {
            FileFormat.TXT: self._extract_txt,
            FileFormat.MD: self._extract_txt,
            FileFormat.DOCX: self._extract_docx,
            FileFormat.PDF: self._extract_pdf,
            FileFormat.RTF: self._extract_rtf,
        }[file_format]

        try:
            text, method_name, warnings = extract_method(file_path)
        except ExtractionError:
            raise
        except Exception as e:
            raise ExtractionError(f"Error inesperado al extraer texto: {str(e)}")

        # Limpiar texto
        text = self._clean_text(text)

        if not text or len(text.strip()) < 20:
            raise ExtractionError(
                "El archivo no contiene suficiente texto útil o está corrupto. "
                "Se requieren al menos ~20 caracteres de contenido."
            )

        # Evaluar calidad
        quality_score = self._evaluate_quality(text)
        word_count = len(text.split())

        return ExtractionResult(
            text=text,
            format_detected=file_format,
            char_count=len(text),
            word_count=word_count,
            estimated_tokens=self._estimate_tokens(text),
            quality_score=round(quality_score, 2),
            extraction_warnings=warnings,
            extraction_method=method_name,
        )

    # --- Parsers específicos ---

    def _extract_txt(self, file_path: Path) -> tuple[str, str, list[str]]:
        """Extrae texto de archivos TXT y MD."""
        warnings = []
        encodings = ["utf-8", "latin-1", "cp1252", "ascii"]

        for encoding in encodings:
            try:
                text = file_path.read_text(encoding=encoding)
                if encoding != "utf-8":
                    warnings.append(f"Archivo leído con encoding {encoding} (no UTF-8)")
                return text, "direct_read", warnings
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ExtractionError(
            "No se pudo leer el archivo de texto. Encoding no reconocido."
        )

    def _extract_docx(self, file_path: Path) -> tuple[str, str, list[str]]:
        """Extrae texto de archivos DOCX incluyendo tablas."""
        warnings = []
        try:
            from docx import Document
        except ImportError:
            raise ExtractionError(
                "La librería python-docx no está instalada. "
                "Ejecuta: pip install python-docx"
            )

        try:
            doc = Document(str(file_path))
            parts = []

            # Extraer párrafos
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Extraer tablas
            for table_idx, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        table_rows.append(" | ".join(cells))
                if table_rows:
                    parts.append(f"\n[Tabla {table_idx + 1}]")
                    parts.extend(table_rows)
                    warnings.append(f"Tabla {table_idx + 1} extraída ({len(table_rows)} filas)")

            return "\n".join(parts), "python-docx", warnings

        except Exception as e:
            raise ExtractionError(f"Error al procesar DOCX: {str(e)}")

    def _extract_pdf(self, file_path: Path) -> tuple[str, str, list[str]]:
        """Extrae texto de archivos PDF (digitales)."""
        warnings = []
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ExtractionError(
                "La librería PyMuPDF no está instalada. "
                "Ejecuta: pip install PyMuPDF"
            )

        try:
            doc = fitz.open(str(file_path))
            parts = []
            pages_without_text = 0

            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text")
                if text and text.strip():
                    parts.append(f"--- Página {page_num} ---\n{text.strip()}")
                else:
                    pages_without_text += 1

            doc.close()

            if pages_without_text > 0:
                warnings.append(
                    f"{pages_without_text} página(s) sin texto detectado "
                    "(posible contenido escaneado/imagen)"
                )

            if not parts:
                raise ExtractionError(
                    "El PDF no contiene texto extraíble. "
                    "Podría ser un documento escaneado. "
                    "Intenta convertirlo a texto primero."
                )

            total_pages = len(list(range(len(doc) if hasattr(doc, '__len__') else 0))) or page_num
            warnings.insert(0, f"PDF procesado: {page_num} páginas")

            return "\n\n".join(parts), "pymupdf", warnings

        except ExtractionError:
            raise
        except Exception as e:
            raise ExtractionError(f"Error al procesar PDF: {str(e)}")

    def _extract_rtf(self, file_path: Path) -> tuple[str, str, list[str]]:
        """Extrae texto de archivos RTF."""
        warnings = []
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ExtractionError(
                "La librería striprtf no está instalada. "
                "Ejecuta: pip install striprtf"
            )

        try:
            raw = file_path.read_text(encoding="utf-8", errors="ignore")
            text = rtf_to_text(raw)
            return text, "striprtf", warnings

        except Exception as e:
            raise ExtractionError(f"Error al procesar RTF: {str(e)}")

    # --- Utilidades ---

    def _clean_text(self, text: str) -> str:
        """Limpia el texto extraído de artefactos comunes."""
        # Remover caracteres de control excepto saltos de línea y tabs
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        # Colapsar múltiples líneas en blanco a máximo 2
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        # Remover espacios trailing
        text = "\n".join(line.rstrip() for line in text.split("\n"))
        return text.strip()

    def _evaluate_quality(self, text: str) -> float:
        """
        Evalúa la calidad del texto en una escala 0-1.
        Factores: ratio de palabras válidas, longitud promedio, caracteres extraños.
        """
        if not text:
            return 0.0

        words = text.split()
        if len(words) < 10:
            return 0.3

        score = 1.0

        # Factor 1: Ratio de caracteres alfanuméricos vs total
        alnum_count = sum(1 for c in text if c.isalnum() or c.isspace())
        alnum_ratio = alnum_count / len(text) if text else 0
        if alnum_ratio < 0.7:
            score -= 0.3

        # Factor 2: Longitud promedio de palabras (palabras muy cortas = OCR roto)
        avg_word_len = sum(len(w) for w in words) / len(words)
        if avg_word_len < 2.5:
            score -= 0.2
        elif avg_word_len > 15:
            score -= 0.1  # Palabras pegadas

        # Factor 3: Ratio de palabras con al menos 3 letras
        real_words = sum(1 for w in words if len(w) >= 3)
        real_word_ratio = real_words / len(words)
        if real_word_ratio < 0.5:
            score -= 0.2

        # Factor 4: Repeticiones excesivas (headers/footers de PDF)
        lines = text.split("\n")
        if len(lines) > 10:
            unique_lines = set(line.strip() for line in lines if line.strip())
            repetition_ratio = len(unique_lines) / len([l for l in lines if l.strip()])
            if repetition_ratio < 0.5:
                score -= 0.2

        return max(0.0, min(1.0, score))

    def _estimate_tokens(self, text: str) -> int:
        """
        Estima tokens de forma aproximada.
        Regla general: ~1 token por cada 4 caracteres en español/inglés.
        """
        return max(1, len(text) // 4)


# Instancia singleton
file_extractor = FileExtractor()
